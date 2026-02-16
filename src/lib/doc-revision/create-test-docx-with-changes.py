#!/usr/bin/env python3
"""
Create a test .docx file with track changes using XML manipulation.
"""

import sys
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn


def add_insertion(paragraph, text, author="Test Author", date=None):
    """Add an insertion (track change) to a paragraph."""
    if date is None:
        date = datetime.now().isoformat()
    
    # Create the insertion element
    ins = OxmlElement('w:ins')
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)
    ins.set(qn('w:id'), '1')
    
    # Create a run inside the insertion
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    ins.append(run)
    
    # Add to paragraph
    paragraph._element.append(ins)


def add_deletion(paragraph, text, author="Test Author", date=None):
    """Add a deletion (track change) to a paragraph."""
    if date is None:
        date = datetime.now().isoformat()
    
    # Create the deletion element
    dels = OxmlElement('w:del')
    dels.set(qn('w:author'), author)
    dels.set(qn('w:date'), date)
    dels.set(qn('w:id'), '2')
    
    # Create a run inside the deletion
    delText = OxmlElement('w:delText')
    delText.text = text
    run = OxmlElement('w:r')
    run.append(delText)
    dels.append(run)
    
    # Add to paragraph
    paragraph._element.append(dels)


def create_test_document(output_path: str):
    """Create a test Word document with track changes."""
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document for Track Changes Extraction', 0)
    
    # Paragraph 1: Normal text
    para1 = doc.add_paragraph('This is the first paragraph with no changes.')
    
    # Paragraph 2: With insertion
    para2 = doc.add_paragraph('This paragraph has an ')
    add_insertion(para2, 'inserted text', author='John Doe', date='2024-01-15T10:30:00Z')
    para2.add_run(' in the middle.')
    
    # Paragraph 3: With deletion
    para3 = doc.add_paragraph('This paragraph ')
    add_deletion(para3, 'had deleted text', author='Jane Smith', date='2024-01-16T14:20:00Z')
    para3.add_run('continues here.')
    
    # Paragraph 4: With both insertion and deletion
    para4 = doc.add_paragraph('This has ')
    add_insertion(para4, 'new content', author='Bob Wilson', date='2024-01-17T09:00:00Z')
    para4.add_run(' and ')
    add_deletion(para4, 'removed content', author='Alice Brown', date='2024-01-17T09:05:00Z')
    para4.add_run(' together.')
    
    # Paragraph 5: Normal ending
    doc.add_paragraph('This is the final paragraph with no changes.')
    
    # Save the document
    doc.save(output_path)
    print(f"Created test document with track changes: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: create-test-docx-with-changes.py <output-path>")
        sys.exit(1)
    
    create_test_document(sys.argv[1])
