#!/usr/bin/env python3
"""
Applies revisions to a Word document while preserving paragraph styles.
Takes revision plan as JSON input and outputs modified .docx file.
"""
import sys
import json
from docx import Document
from docx.shared import Pt, RGBColor
from typing import List, Dict, Any

def apply_revisions_to_word(input_path: str, output_path: str, revision_plan: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Applies revisions to a Word document based on revision plan.
    
    Args:
        input_path: Path to input .docx file
        output_path: Path to save modified .docx file
        revision_plan: List of revisions with paragraph numbers and revised text
        
    Returns:
        List of changes made (change log)
    """
    doc = Document(input_path)
    changes = []
    
    # Process each revision in the plan
    for item in revision_plan:
        para_num = item['paragraphNumber']
        revised_text = item['revisedText']
        feedback_text = item['feedbackText']
        reason = item['reason']
        
        # Validate paragraph number (1-indexed)
        if para_num < 1 or para_num > len(doc.paragraphs):
            continue  # Skip invalid paragraph numbers
        
        para_index = para_num - 1
        paragraph = doc.paragraphs[para_index]
        original_text = paragraph.text
        
        # Only apply change if text actually differs
        if original_text != revised_text:
            # Preserve style
            original_style = paragraph.style
            
            # Clear existing runs and add new text
            paragraph.clear()
            paragraph.add_run(revised_text)
            
            # Restore style
            paragraph.style = original_style
            
            # Record change
            changes.append({
                'lineOrParagraph': para_num,
                'originalText': original_text,
                'revisedText': revised_text,
                'reason': reason,
                'feedbackText': feedback_text
            })
    
    # Save modified document
    doc.save(output_path)
    
    return changes

def main():
    """
    Main entry point. Reads JSON input from stdin and applies revisions.
    
    Expected JSON format:
    {
        "inputPath": "/path/to/input.docx",
        "outputPath": "/path/to/output.docx",
        "revisionPlan": [
            {
                "paragraphNumber": 1,
                "revisedText": "New text",
                "feedbackText": "Original feedback",
                "reason": "Reason for change"
            },
            ...
        ]
    }
    
    Outputs JSON to stdout:
    {
        "success": true,
        "changes": [...]
    }
    """
    try:
        # Read JSON input from stdin
        input_data = json.loads(sys.stdin.read())
        
        input_path = input_data['inputPath']
        output_path = input_data['outputPath']
        revision_plan = input_data['revisionPlan']
        
        # Apply revisions
        changes = apply_revisions_to_word(input_path, output_path, revision_plan)
        
        # Output result as JSON
        result = {
            'success': True,
            'changes': changes
        }
        
        print(json.dumps(result, indent=2))
        sys.exit(0)
        
    except Exception as e:
        error_result = {
            'success': False,
            'error': str(e)
        }
        print(json.dumps(error_result, indent=2))
        sys.exit(1)

if __name__ == '__main__':
    main()
