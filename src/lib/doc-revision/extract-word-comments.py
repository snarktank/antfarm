#!/usr/bin/env python3
"""
Extract comments and track changes from Word documents using python-docx.
"""

import sys
import json
from datetime import datetime
from typing import List, Dict, Any

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print(json.dumps({
        "error": "python-docx not installed. Install with: pip3 install python-docx"
    }))
    sys.exit(1)


def extract_text_from_element(element, include_del_text=False):
    """Extract all text from an XML element recursively."""
    text = ""
    # Get direct text nodes
    for text_elem in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if text_elem.text:
            text += text_elem.text
    
    # Also get deleted text if requested
    if include_del_text:
        for del_text_elem in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}delText'):
            if del_text_elem.text:
                text += del_text_elem.text
    
    return text


def extract_comments(docx_path: str) -> Dict[str, Any]:
    """
    Extract all comments and track changes from a Word document.
    
    Returns a structured dict with:
    - comments: List of comment objects
    - trackChanges: List of tracked change objects
    - paragraphs: List of paragraph texts for context
    """
    try:
        doc = Document(docx_path)
        
        result = {
            "comments": [],
            "trackChanges": [],
            "paragraphs": []
        }
        
        # Extract paragraph texts for context
        for para in doc.paragraphs:
            result["paragraphs"].append(para.text)
        
        # Extract comments from document
        # Comments are stored in the document part
        if hasattr(doc, '_part') and hasattr(doc._part, 'comments_part'):
            comments_part = doc._part.comments_part
            if comments_part:
                comments_element = comments_part.element
                
                for comment in comments_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment'):
                    comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                    date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                    
                    # Extract comment text
                    comment_text = extract_text_from_element(comment)
                    
                    # Find the paragraph that this comment is attached to
                    # This is a simplified approach - in reality, we'd need to traverse the XML more carefully
                    context_para_index = 0
                    
                    result["comments"].append({
                        "id": comment_id,
                        "author": author,
                        "date": date,
                        "text": comment_text.strip(),
                        "paragraphIndex": context_para_index,
                        "contextBefore": [],
                        "contextAfter": []
                    })
        
        # Extract track changes (revisions)
        # Track changes are marked with w:ins (insertions) and w:del (deletions) in the XML
        para_index = 0
        for para in doc.paragraphs:
            para_element = para._element
            
            # Find all insertions in this paragraph
            for ins in para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ins'):
                author = ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                date = ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                text = extract_text_from_element(ins)
                
                if text:  # Only add if there's actual text
                    result["trackChanges"].append({
                        "type": "insertion",
                        "author": author,
                        "date": date,
                        "text": text,
                        "paragraphIndex": para_index
                    })
            
            # Find all deletions in this paragraph
            for dels in para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}del'):
                author = dels.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                date = dels.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                text = extract_text_from_element(dels, include_del_text=True)
                
                if text:  # Only add if there's actual text
                    result["trackChanges"].append({
                        "type": "deletion",
                        "author": author,
                        "date": date,
                        "text": text,
                        "paragraphIndex": para_index
                    })
            
            para_index += 1
        
        return result
        
    except Exception as e:
        return {
            "error": f"Failed to extract comments: {str(e)}"
        }


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(json.dumps({
            "error": "Usage: extract-word-comments.py <path-to-docx>"
        }))
        sys.exit(1)
    
    docx_path = sys.argv[1]
    result = extract_comments(docx_path)
    print(json.dumps(result, indent=2))
