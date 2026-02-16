#!/usr/bin/env python3
"""
Creates a test Word document for revision testing.
"""
from docx import Document
from docx.shared import Pt
import sys

def create_test_revision_docx(output_path: str):
    """Create a simple Word document for testing revisions."""
    doc = Document()
    
    # Add paragraphs with content that will be revised
    p1 = doc.add_paragraph("This is a incorrect fact that needs correction.")
    p2 = doc.add_paragraph("This sentence is missing some important information.")
    p3 = doc.add_paragraph("I'm gonna utilize this wordy sentence in order to demonstrate tone issues.")
    p4 = doc.add_paragraph("This paragraph has good structure and doesn't need changes.")
    
    # Save document
    doc.save(output_path)
    print(f"Created test document: {output_path}")

if __name__ == '__main__':
    output_path = sys.argv[1] if len(sys.argv) > 1 else 'test-revision.docx'
    create_test_revision_docx(output_path)
