#!/usr/bin/env python3
"""
Create a test .docx file with comments for testing the word extractor.
Note: Track changes cannot be easily created programmatically with python-docx.
This creates a document with comments that can be manually edited to add track changes.
"""

import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def create_test_document(output_path: str):
    """Create a test Word document with some paragraphs."""
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document for Comment and Track Changes Extraction', 0)
    
    # Add some paragraphs
    doc.add_paragraph('This is the first paragraph. It contains some sample text.')
    doc.add_paragraph('This is the second paragraph. It has more text for testing purposes.')
    doc.add_paragraph('This is the third paragraph. We will add comments and track changes to this document.')
    doc.add_paragraph('This is the fourth paragraph. The reader agent will extract all annotations.')
    doc.add_paragraph('This is the fifth and final paragraph. Testing is important!')
    
    # Save the document
    doc.save(output_path)
    print(f"Created test document: {output_path}")
    print("To add comments and track changes:")
    print("  1. Open the document in Microsoft Word")
    print("  2. Add some comments (Insert > Comment)")
    print("  3. Enable Track Changes (Review > Track Changes)")
    print("  4. Make some edits (insertions and deletions)")
    print("  5. Save the document")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: create-test-docx.py <output-path>")
        sys.exit(1)
    
    create_test_document(sys.argv[1])
